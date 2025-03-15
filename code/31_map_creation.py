import pandas as pd
from pathlib import Path
import matplotlib.pyplot as plt
import geopandas as gpd
from docx import Document
from docx.shared import Inches

home = Path.home()
work_dir = (home / 'medicaid_project')
data = (work_dir / 'data')
shapefiles = (data / 'shapefiles')
raw_data = (data / 'raw')
clean_data = (data / 'clean')
state_level = (clean_data / 'state_level')
output = (work_dir / 'output')
maps = (output / 'maps')
code = Path.cwd() 

## pull in county-level and state-level master data 
state_level = pd.read_csv(f'{clean_data}/master_state_level.csv')
# our state_level dataste includes territories; we don't want this. 
state_level = state_level[state_level['usa_state_dummy'] == 1]
state_level = state_level[state_level['state'] != 'alaska']
state_level = state_level[state_level['state'] != 'hawaii']

## pull in USA state-level shapefiles
path = f'{shapefiles}/tl_2024_us_state/tl_2024_us_state.shp'
states_gdf = gpd.read_file(path)
states_gdf["NAME"] = states_gdf["NAME"].str.lower()
states_gdf = states_gdf[~states_gdf['NAME'].isin([
    'united states virgin islands', 
    'commonwealth of the northern mariana islands',
    'guam', 
    'american samoa', 
    'puerto rico',
])]

doc = Document()
doc.add_heading('Medicaid Enrollment Maps', level=1)

######################################################################################################################################
# State-Level Map of Medicaid Enrollment in 2024
######################################################################################################################################
'''for age, title in zip(['', '_under_19', '_19_to_64', '_65_and_over'],
                      ['', '- Under 19', '- 19-64', '- 65 and Over']):'''
    #restrict to 2024
for age, title in zip([''], ['']):
    df_2024 = state_level[state_level['year'] == 2023]

    merged_gdf = states_gdf.merge(df_2024, left_on='NAME', right_on='state', how='left')

    fig, ax = plt.subplots(figsize=(12, 8))
    merged_gdf.plot(
        column=f'pct_enrollment{age}_medicaid_chip',    # difference column to color by
        cmap='OrRd',          # color map (choose any you like, e.g., 'YlGnBu', 'viridis')
        legend=True,          # show the color scale
        edgecolor='black',    # outline for states
        linewidth=0.4,
        ax=ax,
        legend_kwds={'shrink': 0.5}
    )
    # Add a title and remove axis clutter
    ax.set_title(f'Medicaid Enrollment (%) (2023) {title}')
    ax.set_axis_off()
    image_path = maps / f'medicaid_enrollment_2023{age}.png'
    plt.savefig(image_path, dpi=300, bbox_inches='tight')
    plt.show()

    # Add to Word Document
    doc.add_paragraph(title)
    doc.add_picture(str(image_path), width=Inches(6))

######################################################################################################################################
# State-Level Map of Where Greatest Change in Medicaid Enrollment
######################################################################################################################################
'''for age, title in zip(['', '_under_19', '_19_to_64', '_65_and_over'],
                      ['', '- Under 19', '- 19-64', '- 65 and Over']):'''
for age, title in zip([''], ['']):
    # first reshape dataframe wide 
    wide = state_level.pivot(
    index='state', 
    columns='year', 
    values=f'pct_enrollment{age}_medicaid_chip'
    ).reset_index()
    # Rename the pivoted columns for clarity:
    wide.rename(columns={2017: 'pct_2017', 2023: 'pct_2023'}, inplace=True)
    # Compute difference (may produce NaN if data is missing for a given state/year)
    wide['pct_diff'] = wide['pct_2023'] - wide['pct_2017']
    # merge data with shapefile
    merged_gdf = states_gdf.merge(wide, left_on='NAME', right_on='state', how='left')

    fig, ax = plt.subplots(figsize=(12, 8))
    merged_gdf.plot(
        column='pct_diff',    # difference column to color by
        cmap='OrRd',          # color map (choose any you like, e.g., 'YlGnBu', 'viridis')
        legend=True,          # show the color scale
        edgecolor='black',    # outline for states
        linewidth=0.4,
        ax=ax,
        legend_kwds={'shrink': 0.5}
    )
    # Add a title and remove axis clutter
    ax.set_title(f'Change in Medicaid Enrollment (%) from 2017 to 2023 {title}')
    ax.set_axis_off()
    image_path = maps / f'medicaid_enrollment_change{age}.png'
    plt.savefig(image_path, dpi=300, bbox_inches='tight')
    plt.show()

    # Add to Word Document
    doc.add_paragraph(title)
    doc.add_picture(str(image_path), width=Inches(6))

doc_path = maps / 'medicaid_maps.docx'
doc.save(doc_path)
print(f"✅ Word document saved: {doc_path}")